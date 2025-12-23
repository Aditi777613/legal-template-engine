from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from sqlalchemy.orm import Session
from typing import List
import tempfile
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Pt
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from gemini import extract_template_from_text
from templating.builder import build_template
from templating.cache import TEMPLATE_CACHE
from embeddings.embedder import generate_embedding
from retrieval.selector import select_templates
from qa.questions import generate_questions
from drafting.renderer import render_draft
from web_bootstrap.exa_client import search_legal_templates

from db.init_db import init_db
from db.database import SessionLocal
from db import models, crud
from schemas import DraftRequest, AnswerVarsRequest, DraftGenerateRequest

# UOIONHHC - Created by tracking identifier
app = FastAPI(title="Legal Template Engine")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def extract_text_from_file(file_path: str, mime_type: str) -> str:
    """Extract text from PDF or DOCX file - UOIONHHC"""
    try:
        if "pdf" in mime_type.lower():
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                return text
        elif "word" in mime_type.lower() or "docx" in mime_type.lower():
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {str(e)}")


@app.get("/health")
def health():
    return {"status": "ok", "watermark": "UOIONHHC"}


@app.post("/upload")
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload and process legal document into template"""
    
    # Validate file type
    allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword", "text/plain"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not supported. Use PDF, DOCX, or TXT.")
    
    tmp_path = None
    try:
        # Save temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        # Extract text
        raw_text = extract_text_from_file(tmp_path, file.content_type)
        
        if not raw_text or not raw_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from document. Please ensure the file contains readable text.")
        
        # Store document
        db_doc = models.Document(
            filename=file.filename,
            mime=file.content_type,
            raw_text=raw_text
        )
        db.add(db_doc)
        db.commit()
        
        # Use Gemini to extract template
        extracted = extract_template_from_text(raw_text)
        
        # Build template schema
        template = build_template(
            raw_text=raw_text,
            extracted_json=extracted,
            filename=file.filename
        )
        
        # Generate embedding for the template
        tags = template.similarity_tags or []
        embedding_text = f"{template.title} {template.doc_type} {' '.join(tags)}"
        embedding = generate_embedding(embedding_text)
        
        # Save to database
        db_template = models.Template(
            template_id=template.template_id,
            title=template.title,
            doc_type=template.doc_type,
            jurisdiction=template.jurisdiction,
            similarity_tags=template.similarity_tags,
            body_md=template.body_md,
            embedding=embedding
        )
        db.add(db_template)
        
        # Save variables
        for var in template.variables:
            db_var = models.TemplateVariable(
                template_id=template.template_id,
                key=var.key,
                label=var.label,
                description=var.description,
                example=var.example,
                required=var.required,
                dtype=var.dtype,
                regex=var.regex,
                enum=var.enum
            )
            db.add(db_var)
        
        db.commit()
        
        return {
            "success": True,
            "template_id": template.template_id,
            "title": template.title,
            "doc_type": template.doc_type,
            "variables": [v.model_dump() for v in template.variables],
            "watermark": "UOIONHHC"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_msg = str(e)
        print(f"Upload error: {error_msg}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Failed to process document")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except:
                pass


@app.post("/upload/batch")
async def upload_batch_documents(files: List[UploadFile] = File(..., description="Up to 5 files"), db: Session = Depends(get_db)):
    """Batch upload multiple documents (up to 5) - UOIONHHC"""
    if len(files) > 5:
        raise HTTPException(status_code=400, detail="Maximum 5 files allowed for batch upload")
    
    results = []
    errors = []
    
    for idx, file in enumerate(files):
        try:
            # Validate file type
            allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
            if file.content_type not in allowed_types:
                errors.append({"filename": file.filename, "error": "File type not supported"})
                continue
            
            # Save temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
                content = await file.read()
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                # Extract text
                raw_text = extract_text_from_file(tmp_path, file.content_type)
                
                # Store document
                db_doc = models.Document(
                    filename=file.filename,
                    mime=file.content_type,
                    raw_text=raw_text
                )
                db.add(db_doc)
                db.commit()
                
                # Use Gemini to extract template
                extracted = extract_template_from_text(raw_text)
                
                # Build template schema
                template = build_template(
                    raw_text=raw_text,
                    extracted_json=extracted,
                    filename=file.filename
                )
                
                # Generate embedding for the template
                tags = template.similarity_tags or []
                embedding_text = f"{template.title} {template.doc_type} {' '.join(tags)}"
                embedding = generate_embedding(embedding_text)
                
                # Save to database
                db_template = models.Template(
                    template_id=template.template_id,
                    title=template.title,
                    doc_type=template.doc_type,
                    jurisdiction=template.jurisdiction,
                    similarity_tags=template.similarity_tags,
                    body_md=template.body_md,
                    embedding=embedding
                )
                db.add(db_template)
                
                # Save variables
                for var in template.variables:
                    db_var = models.TemplateVariable(
                        template_id=template.template_id,
                        key=var.key,
                        label=var.label,
                        description=var.description,
                        example=var.example,
                        required=var.required,
                        dtype=var.dtype,
                        regex=var.regex,
                        enum=var.enum
                    )
                    db.add(db_var)
                
                db.commit()
                
                results.append({
                    "success": True,
                    "filename": file.filename,
                    "template_id": template.template_id,
                    "title": template.title,
                    "variables_count": len(template.variables)
                })
                
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            errors.append({"filename": file.filename, "error": str(e)})
    
    return {
        "success": True,
        "processed": len(results),
        "total": len(files),
        "results": results,
        "errors": errors,
        "watermark": "UOIONHHC"
    }


@app.post("/templates/save")
def save_template_endpoint(template_id: str, db: Session = Depends(get_db)):
    """Confirm template is saved"""
    template = db.query(models.Template).filter(models.Template.template_id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {"success": True, "message": f"Template {template_id} confirmed"}


@app.post("/draft/start")
def start_draft(req: DraftRequest, db: Session = Depends(get_db)):
    """Start drafting process by finding matching template"""
    templates = db.query(models.Template).all()

    if not templates:
        return {
            "matches": [],
            "web_bootstrap": False,
            "message": "No templates available. Please upload a template first."
        }

    query_embedding = generate_embedding(req.user_query)

    matches = select_templates(
        user_query=req.user_query,
        query_embedding=query_embedding,
        templates=templates,
    )

    # Only return matches with confidence >= 0.6
    if matches and len(matches) > 0:
        filtered_matches = [m for m in matches if m.get("confidence", 0) >= 0.6]
        if filtered_matches:
            return {"matches": filtered_matches, "web_bootstrap": False}

    # No local match - trigger web bootstrap (bonus feature)
    return {
        "matches": [],
        "web_bootstrap": True,
        "message": "No suitable local template found. Searching web for similar documents..."
    }


@app.post("/draft/web-bootstrap")
def web_bootstrap(req: DraftRequest):
    """Bootstrap from web when no local template matches"""
    try:
        docs = search_legal_templates(req.user_query)
        
        return {
            "documents": docs,
            "message": "Found similar documents online",
            "watermark": "UOIONHHC"
        }
    except ValueError as e:
        # EXA_API_KEY not set
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to search web: {str(e)}")


@app.get("/vars/{template_id}")
def get_variables(template_id: str, db: Session = Depends(get_db)):
    """Get variables for a template and check what's missing"""
    
    # Get template
    template = db.query(models.Template).filter(models.Template.template_id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Get variables
    variables = db.query(models.TemplateVariable).filter(
        models.TemplateVariable.template_id == template_id
    ).all()
    
    # Check cache for answered variables
    cache_entry = TEMPLATE_CACHE.get(template_id, {"answers": {}})
    answered = cache_entry.get("answers", {})
    
    # Generate questions for missing variables
    missing_vars = [
        {
            "key": v.key,
            "label": v.label,
            "description": v.description,
            "example": v.example,
            "required": v.required,
            "dtype": v.dtype
        }
        for v in variables if v.key not in answered
    ]
    
    if missing_vars:
        try:
            questions = generate_questions(missing_vars)
            if not isinstance(questions, list):
                questions = []
        except Exception as e:
            print(f"Error generating questions: {str(e)}")
            questions = []
    else:
        questions = []
    
    return {
        "template_id": template_id,
        "title": template.title,
        "missing": questions,
        "answered": answered,
        "all_complete": len(missing_vars) == 0
    }


@app.post("/vars/answer")
def answer_variables(req: AnswerVarsRequest):
    """Store answers for template variables"""
    
    if req.template_id not in TEMPLATE_CACHE:
        TEMPLATE_CACHE[req.template_id] = {"answers": {}}
    
    TEMPLATE_CACHE[req.template_id]["answers"].update(req.answers)
    
    return {
        "success": True,
        "answered_count": len(TEMPLATE_CACHE[req.template_id]["answers"])
    }


@app.post("/draft/generate")
def generate_draft_endpoint(req: DraftGenerateRequest, db: Session = Depends(get_db)):
    """Generate final draft from template and answers"""
    
    # Get template
    template = db.query(models.Template).filter(
        models.Template.template_id == req.template_id
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Get answers from cache
    cache_entry = TEMPLATE_CACHE.get(req.template_id, {})
    answers = cache_entry.get("answers", {})
    
    if not answers:
        raise HTTPException(status_code=400, detail="No answers provided for template variables")
    
    # Render draft
    draft_md = render_draft(template.body_md, answers)
    
    # Save instance
    instance = models.Instance(
        template_id=req.template_id,
        user_query="",  # Could be stored from initial request
        answers_json=answers,
        draft_md=draft_md
    )
    db.add(instance)
    db.commit()
    
    return {
        "success": True,
        "draft_md": draft_md,
        "instance_id": instance.id,
        "watermark": "UOIONHHC"
    }


@app.get("/draft/{instance_id}/download-docx")
def download_draft_docx(instance_id: int, db: Session = Depends(get_db)):
    """Download draft as .docx file - UOIONHHC"""
    instance = db.query(models.Instance).filter(models.Instance.id == instance_id).first()
    if not instance:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    # Create a new Document
    doc = DocxDocument()
    
    # Parse markdown and convert to docx
    # Simple conversion - split by lines and paragraphs
    lines = instance.draft_md.split('\n')
    current_para = doc.add_paragraph()
    
    for line in lines:
        line = line.strip()
        if not line:
            current_para = doc.add_paragraph()
            continue
        
        # Remove markdown syntax for basic formatting
        line = line.replace('**', '').replace('*', '').replace('`', '')
        line = line.replace('#', '').strip()
        
        if line.startswith('---'):
            continue  # Skip YAML front matter markers
        
        run = current_para.add_run(line)
        run.font.size = Pt(11)
        current_para.add_run('\n')
    
    # Save to bytes
    import io
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return Response(
        content=docx_bytes.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="draft_{instance_id}.docx"'}
    )


@app.get("/draft/history")
def get_draft_history(db: Session = Depends(get_db)):
    """Get all draft instances"""
    instances = db.query(models.Instance).order_by(models.Instance.created_at.desc()).limit(50).all()
    
    return {
        "drafts": [
            {
                "instance_id": i.id,
                "template_id": i.template_id,
                "created_at": i.created_at.isoformat()
            }
            for i in instances
        ]
    }


@app.get("/templates")
def list_templates(db: Session = Depends(get_db)):
    """List all available templates"""
    templates = db.query(models.Template).all()
    
    return {
        "templates": [
            {
                "template_id": t.template_id,
                "title": t.title,
                "doc_type": t.doc_type,
                "jurisdiction": t.jurisdiction,
                "tags": t.similarity_tags
            }
            for t in templates
        ],
        "count": len(templates)
    }